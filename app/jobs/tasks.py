from datetime import datetime, timezone
from pathlib import Path
import csv
import io
import re
import subprocess
import tempfile

from sqlalchemy import select
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader
import xlrd
from docx import Document
from PIL import Image
import pytesseract
import pdfplumber

from app.db.session import SessionLocal
from app.core.config import settings
from app.models import ApprovedMatch, ExportFile, MatchResult, ParsedItem, ProductCatalog, RequestItem, UploadedFile
from app.services.catalog_match_rules import apply_catalog_rule_filter, list_catalog_match_rules
from app.services.hybrid_search import ensure_catalog_index, vector_search
from app.services.stop_words import load_stop_words
from app.services.storage import export_path


_STOP_WORDS = {
    "для",
    "под",
    "над",
    "без",
    "с",
    "со",
    "и",
    "в",
    "на",
    "к",
    "по",
    "шт",
    "комплект",
    "метр",
    "метров",
    "метра",
    "штук",
}


def _normalize_text(value: str | None) -> str:
    if not value:
        return ""
    text = value.lower().replace("ё", "е")
    text = re.sub(r"[^a-zа-я0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _tokenize(value: str | None) -> set[str]:
    text = _normalize_text(value)
    if not text:
        return set()
    norm_map = {
        "присоединителей": "присоединитель",
        "присоединителя": "присоединитель",
        "присоединители": "присоединитель",
        "присоединит": "присоединитель",
        "радиаторам": "радиатор",
        "радиаторов": "радиатор",
        "счетчика": "счетчик",
        "счетчики": "счетчик",
        "ре": "pe",
        "ха": "xa",
        "реха": "pexa",
        "pexa": "pexa",
    }
    keep_short = {"pe", "xa", "du", "ре", "ха"}
    tokens = []
    for t in text.split():
        if t.isdigit():
            continue
        if (len(t) < 3 and t not in keep_short) or t in _STOP_WORDS:
            continue
        tokens.append(norm_map.get(t, t))
    return set(tokens)


def _name_similarity(item_name: str | None, product_name: str | None) -> float:
    item_norm = _normalize_text(item_name)
    product_norm = _normalize_text(product_name)
    if not item_norm or not product_norm:
        return 0.0
    if item_norm in product_norm or product_norm in item_norm:
        return 0.85
    item_tokens = _tokenize(item_norm)
    product_tokens = _tokenize(product_norm)
    if not item_tokens or not product_tokens:
        return 0.0
    overlap = len(item_tokens & product_tokens)
    if overlap == 0:
        return 0.0
    recall = overlap / len(item_tokens)
    precision = overlap / len(product_tokens)
    important_roots = {"присоединитель", "счетчик", "распределитель", "радиатор"}
    root_bonus = 0.0
    if any(root in item_tokens and root in product_tokens for root in important_roots):
        root_bonus = 0.15
    return round(min(0.95, max(recall, precision) * 0.8 + root_bonus), 4)


def _size_tokens(value: str | None) -> set[str]:
    text = _normalize_text(value)
    if not text:
        return set()
    tokens = set(re.findall(r"\b\d+/\d+\b", text))
    tokens.update(re.findall(r"\bdu\s*\d+\b", text))
    tokens.update(re.findall(r"\bду\s*\d+\b", text))
    tokens.update(re.findall(r"\b\d+(?:[.,]\d+)?\s*мм\b", text))
    # compact normalization (du 15 -> du15)
    compact = set()
    for t in tokens:
        compact.add(re.sub(r"\s+", "", t.replace(",", ".")))
    return compact


def _number_tokens(value: str | None) -> set[str]:
    text = _normalize_text(value)
    if not text:
        return set()
    nums = set(re.findall(r"\b\d+(?:[.,]\d+)?(?:/\d+(?:[.,]\d+)?)?\b", text))
    return {n.replace(",", ".") for n in nums}


def _du_tokens(value: str | None) -> set[str]:
    text = _normalize_text(value)
    if not text:
        return set()
    found = re.findall(r"\b(?:ду|du)\s*(\d+)\b", text)
    return {f"du{x}" for x in found}


def _fraction_tokens(value: str | None) -> set[str]:
    text = _normalize_text(value)
    if not text:
        return set()
    return set(re.findall(r"\b\d+/\d+\b", text))


def _semantic_query_text(value: str | None) -> str:
    text = _normalize_text(value)
    if not text:
        return ""
    text = re.sub(r"\([^)]*\)", " ", text)
    text = re.sub(r"\b\d{3,}\b", " ", text)
    text = re.sub(r"\b(метр|метров|метра|шт|штук|упаковка|уп)\b", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _expanded_queries(value: str | None) -> list[str]:
    base = _semantic_query_text(value)
    if not base:
        return []
    variants = {base}
    alt = base.replace("ре ха", "pe xa").replace("ре-ха", "pe xa")
    variants.add(alt)
    variants.add(alt.replace("pe xa", "pexa"))
    return [x for x in variants if x]


def _vector_recall_scores(query: str | None) -> dict[str, float]:
    if not query:
        return {}
    queries = _expanded_queries(query)
    queries.extend([q for q in [_normalize_text(query)] if q])
    merged: dict[str, float] = {}
    for q in queries:
        try:
            scores = vector_search(q, top_k=80)
        except Exception:
            continue
        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        for rank, (pid, score) in enumerate(ranked, start=1):
            rrf = 1.0 / (60 + rank)
            prev = merged.get(pid, 0.0)
            merged[pid] = max(prev, float(score) * 0.85 + rrf * 2.5)
    return {k: max(0.0, min(1.0, v)) for k, v in merged.items()}


def _dimension_number_tokens(value: str | None) -> set[str]:
    if not value:
        return set()
    text = value.lower().replace("х", "x")
    found = re.findall(r"\b\d{1,3}(?:\s*x\s*\d{1,3}){1,2}\b", text)
    out: set[str] = set()
    for chunk in found:
        out.update(re.findall(r"\d{1,3}", chunk))
    return out


def _spec_number_tokens(value: str | None) -> set[str]:
    tokens: set[str] = set()
    tokens.update(_dimension_number_tokens(value))
    tokens.update(_du_tokens(value))
    tokens.update(_fraction_tokens(value))
    for raw in _number_tokens(value):
        if "/" in raw:
            tokens.add(raw)
            continue
        try:
            v = float(raw)
        except Exception:
            continue
        # Typical product specs (diameter/size). Ignore big quantities from КП.
        if 1 <= v <= 125:
            tokens.add(str(int(v)) if v.is_integer() else str(v))
    return tokens


def _semantic_penalty(item_name: str | None, product_name: str | None) -> tuple[float, str | None]:
    item = _normalize_text(item_name)
    product = _normalize_text(product_name)
    if not item or not product:
        return 0.0, None
    if "труба" in item:
        if "труба" not in product:
            return 0.28, "semantic_penalty:not_pipe_product"
        accessory = (
            "фиксатор",
            "гильза",
            "муфта",
            "тройник",
            "угольник",
            "скоба",
            "инструм",
            "ножниц",
            "накладк",
            "переход",
            "заглушк",
            "шина",
        )
        if any(x in product for x in accessory):
            return 0.22, "semantic_penalty:tube_accessory"
    if "теплоизоляц" in item and ("теплоизоляц" not in product and "изоляц" not in product):
        return 0.2, "semantic_penalty:insulation_missing"
    if "кронштейн" in item and ("кроншт" not in product and "водорозет" not in product):
        return 0.15, "semantic_penalty:bracket_mismatch"
    return 0.0, None


def _semantic_bonus(item_name: str | None, product_name: str | None) -> tuple[float, str | None]:
    item = _normalize_text(item_name)
    product = _normalize_text(product_name)
    if not item or not product:
        return 0.0, None
    if "труба" in item and ("pe xa" in item or "pexa" in item or "ре ха" in item):
        if "труба" in product and ("pe xa" in product or "pexa" in product):
            return 0.24, "semantic_bonus:tube_pex"
    if "кронштейн" in item and "водорозет" in item:
        if "водорозет" in product:
            return 0.2, "semantic_bonus:water_socket"
    if "теплоизоляц" in item and ("теплоизоляц" in product or "изоляц" in product):
        return 0.18, "semantic_bonus:insulation"
    return 0.0, None


def _token_overlap_bonus(item_name: str | None, product_name: str | None) -> tuple[float, str | None]:
    item_tokens = _tokenize(item_name)
    product_tokens = _tokenize(product_name)
    if not item_tokens or not product_tokens:
        return 0.0, None
    generic = {"латунь", "аксиальный", "аналог", "тип", "для", "sanext"}
    signal = [t for t in item_tokens if t not in generic]
    if not signal:
        return 0.0, None
    overlap = len(set(signal) & product_tokens) / max(1, len(set(signal)))
    if overlap >= 0.75:
        return 0.22, f"semantic_bonus:token_overlap_{round(overlap,2)}"
    if overlap >= 0.5:
        return 0.12, f"semantic_bonus:token_overlap_{round(overlap,2)}"
    return 0.0, None


def _product_classes(value: str | None) -> set[str]:
    text = _normalize_text(value)
    tokens = _tokenize(value)
    classes: set[str] = set()
    if any(t.startswith("счетчик") for t in tokens):
        classes.add("meter")
    if any(t.startswith("присоединитель") for t in tokens):
        classes.add("connector")
    if any(t.startswith("распределитель") for t in tokens):
        classes.add("distributor")
    if any(t.startswith("радиатор") for t in tokens):
        classes.add("radiator")
    if any(t in {"комплект", "набор", "монтажный"} for t in tokens):
        classes.add("kit")
    # "счетчик ... с присоединителями" should remain a meter position, not connector.
    if "meter" in classes and "connector" in classes and re.search(r"\bс\s+присоединител", text):
        classes.discard("connector")
    return classes


def _primary_class(value: str | None) -> str | None:
    text = _normalize_text(value)
    if not text:
        return None
    # "комплект присоединителей" is still a connector position, not generic kit.
    if re.search(r"\bкомплект\b", text) and re.search(r"\bприсоединител", text):
        return "connector"
    if re.search(r"\bкомплект\b", text) and re.search(r"\bсчетчик", text):
        return "meter"
    patterns = [
        ("meter", r"\bсчетчик"),
        ("connector", r"\bприсоединител"),
        ("distributor", r"\bраспределител"),
        ("radiator", r"\bрадиатор"),
        ("kit", r"\bкомплект|\bнабор|\bмонтажн"),
    ]
    best: tuple[int, str] | None = None
    for klass, pattern in patterns:
        m = re.search(pattern, text)
        if not m:
            continue
        if best is None or m.start() < best[0]:
            best = (m.start(), klass)
    return best[1] if best else None


def _class_present_in_name(value: str | None, klass: str) -> bool:
    text = _normalize_text(value)
    if not text:
        return False
    patterns = {
        "meter": r"\bсчетчик",
        "connector": r"\bприсоединител|\bвставк",
        "distributor": r"\bраспределител",
        "radiator": r"\bрадиатор",
        "kit": r"\bкомплект|\bнабор|\bмонтажн",
    }
    pattern = patterns.get(klass)
    return bool(pattern and re.search(pattern, text))


def _contains_stop_word(item_name: str | None, stop_words: set[str]) -> str | None:
    text = _normalize_text(item_name)
    if not text:
        return None
    tokens = set(text.split())
    for w in stop_words:
        if not w:
            continue
        norm = _normalize_text(w)
        if not norm:
            continue
        if norm in text or norm in tokens:
            return w
    return None


def _first_non_empty(row: dict, aliases: list[str]) -> str | None:
    for key in aliases:
        value = row.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return None


def _extract_rows_from_file(storage_key: str) -> list[dict]:
    path = Path(storage_key)
    if not path.exists():
        return []
    lower = path.name.lower()
    if lower.endswith(".csv"):
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
        return [dict(r) for r in csv.DictReader(io.StringIO(text))]
    if lower.endswith(".xlsx") or lower.endswith(".xlsm"):
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        values = list(ws.iter_rows(values_only=True))
        if not values:
            return []
        headers = [str(h).strip().lower() if h else "" for h in values[0]]
        result: list[dict] = []
        for row in values[1:]:
            item = {}
            for idx, value in enumerate(row):
                key = headers[idx] if idx < len(headers) else f"col_{idx + 1}"
                item[key] = value
            if any(v is not None and str(v).strip() for v in item.values()):
                result.append(item)
        return result
    if lower.endswith(".xls"):
        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        if ws.nrows == 0:
            return []
        headers = [str(ws.cell_value(0, c)).strip().lower() if ws.cell_value(0, c) else "" for c in range(ws.ncols)]
        result: list[dict] = []
        for r in range(1, ws.nrows):
            item: dict = {}
            for c in range(ws.ncols):
                key = headers[c] if c < len(headers) else f"col_{c + 1}"
                item[key] = ws.cell_value(r, c)
            if any(v is not None and str(v).strip() for v in item.values()):
                result.append(item)
        return result
    if lower.endswith(".pdf"):
        rows: list[dict] = []
        # 1) Prefer structured table extraction for invoices/specs.
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages[:10]:
                    for table in page.extract_tables() or []:
                        if not table or len(table) < 2:
                            continue
                        header = [str(c or "").strip().lower() for c in table[0]]
                        for raw in table[1:]:
                            cells = [str(c or "").strip() for c in raw]
                            if not any(cells):
                                continue
                            normalized = {header[i] if i < len(header) else f"col_{i+1}": cells[i] for i in range(len(cells))}
                            name = _first_non_empty(
                                normalized,
                                ["товары (работы, услуги)", "товары", "наименование", "позиция", "item", "name"],
                            )
                            qty = _first_non_empty(normalized, ["кол-во", "количество", "qty", "quantity"])
                            unit = _first_non_empty(normalized, ["ед.", "ед", "unit"])
                            if name:
                                rows.append(
                                    {
                                        "name": name,
                                        "quantity": qty or "1",
                                        "unit": unit or "",
                                        "article": _first_non_empty(normalized, ["артикул", "sku", "код"]),
                                        "brand": _first_non_empty(normalized, ["бренд", "brand", "производитель"]),
                                    }
                                )
        except Exception:
            pass
        if rows:
            return rows

        # 2) Fallback to plain text extraction and row-pattern parsing.
        reader = PdfReader(str(path))
        lines: list[dict] = []
        for page in reader.pages[:5]:
            text = (page.extract_text() or "").splitlines()
            if not text or len("".join(text).strip()) < 20:
                # OCR fallback for scanned PDFs (optional, depends on local tesseract installation).
                for img in getattr(page, "images", []):
                    try:
                        image = Image.open(io.BytesIO(img.data))
                        ocr_text = pytesseract.image_to_string(image, lang="rus+eng")
                        text.extend([x for x in ocr_text.splitlines() if x.strip()])
                    except Exception:
                        continue
            for line in text:
                line = line.strip()
                if line:
                    # Typical row: "1 Наименование ... 360 шт"
                    m = re.match(r"^\s*\d+\s+(.+?)\s+(\d+(?:[.,]\d+)?)\s+([A-Za-zА-Яа-я]+)\s*$", line)
                    if m:
                        lines.append(
                            {
                                "name": m.group(1).strip(),
                                "quantity": m.group(2).replace(",", "."),
                                "unit": m.group(3).strip(),
                            }
                        )
                        continue
                    cols = [c.strip() for c in re.split(r"\t+| {2,}|;", line) if c.strip()]
                    if len(cols) >= 4:
                        lines.append(
                            {
                                "name": cols[0],
                                "article": cols[1],
                                "brand": cols[2],
                                "quantity": cols[3],
                            }
                        )
                    else:
                        lines.append({"name": line})
        # remove obvious service lines that are not positions
        filtered = []
        noise = {"итого", "всего", "покупатель", "поставщик", "бухгалтер", "руководитель"}
        for item in lines:
            text = (item.get("name") or "").lower()
            if any(n in text for n in noise):
                continue
            filtered.append(item)
        return filtered
    if lower.endswith(".docx"):
        doc = Document(str(path))
        rows: list[dict] = []
        for table in doc.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                item: dict = {}
                for idx, cell in enumerate(row.cells):
                    key = headers[idx] if idx < len(headers) and headers[idx] else f"col_{idx + 1}"
                    item[key] = cell.text.strip()
                if any(str(v).strip() for v in item.values()):
                    rows.append(item)
        if rows:
            return rows
        lines = [{"name": p.text.strip()} for p in doc.paragraphs if p.text and p.text.strip()]
        return lines
    if lower.endswith(".doc"):
        # Best-effort parsing for legacy DOC.
        try:
            out = subprocess.run(["antiword", str(path)], capture_output=True, text=True, timeout=20)
            if out.returncode == 0 and out.stdout.strip():
                return [{"name": line.strip()} for line in out.stdout.splitlines() if line.strip()]
        except Exception:
            pass
        try:
            with tempfile.TemporaryDirectory() as td:
                conv = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", td, str(path)],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if conv.returncode == 0:
                    txt_files = list(Path(td).glob("*.txt"))
                    if txt_files:
                        text = txt_files[0].read_text(encoding="utf-8", errors="ignore")
                        return [{"name": line.strip()} for line in text.splitlines() if line.strip()]
        except Exception:
            pass
        return [{"name": f"legacy doc file: {path.name}"}]
    if lower.endswith(".txt"):
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [{"name": line.strip()} for line in text.splitlines() if line.strip()]
    return []


def _rows_to_parsed_lines(rows: list[dict]) -> list[dict]:
    parsed: list[dict] = []
    for row in rows:
        normalized = {str(k).strip().lower(): v for k, v in row.items()}
        item_name = _first_non_empty(normalized, ["name", "наименование", "товар", "позиция", "item_name"])
        article = _first_non_empty(normalized, ["article", "артикул", "sku", "код"])
        brand = _first_non_empty(normalized, ["brand", "бренд", "производитель"])
        qty_val = _first_non_empty(normalized, ["qty", "quantity", "количество", "кол-во"])
        try:
            quantity = float((qty_val or "1").replace(",", "."))
        except ValueError:
            quantity = 1.0
        if not item_name:
            item_name = article or "unparsed item"
        raw_text = " | ".join(
            [str(x) for x in [item_name, article or "", brand or "", quantity] if str(x).strip()]
        )
        parsed.append(
            {
                "raw_text": raw_text,
                "item_name": item_name,
                "article": article,
                "brand": brand,
                "quantity": quantity,
                "specs": {"source_row": normalized},
                "parse_confidence": 0.9 if item_name else 0.6,
            }
        )
    return parsed


def parsing_task(request_id: str) -> None:
    with SessionLocal() as db:
        req = db.get(RequestItem, request_id)
        if not req:
            return
        req.parse_status = "running"
        db.commit()
        parsed_rows: list[dict] = []
        if req.input_text:
            parsed_rows = _rows_to_parsed_lines([{"name": line.strip()} for line in req.input_text.splitlines() if line.strip()])
        if not parsed_rows:
            latest_file = db.scalar(select(UploadedFile).where(UploadedFile.request_id == request_id).order_by(UploadedFile.created_at.desc()))
            if latest_file:
                rows = _extract_rows_from_file(latest_file.storage_key)
                parsed_rows = _rows_to_parsed_lines(rows)
        if not parsed_rows:
            parsed_rows = _rows_to_parsed_lines([{"name": "Насос Grundfos CR 3-15, 2 шт", "article": "CR3-15", "brand": "Grundfos", "quantity": 2}])
        db.query(ParsedItem).filter(ParsedItem.request_id == request_id).delete()
        for row in parsed_rows:
            db.add(
                ParsedItem(
                    request_id=request_id,
                    raw_text=row["raw_text"],
                    item_name=row["item_name"],
                    article=row["article"],
                    brand=row["brand"],
                    quantity=row["quantity"],
                    parse_confidence=row["parse_confidence"],
                    specs=row["specs"],
                )
            )
        req.total_items = len(parsed_rows)
        req.parse_status = "completed"
        req.status = "parsed"
        req.updated_at = datetime.now(timezone.utc)
        db.commit()


def matching_task(request_id: str, threshold: float) -> None:
    with SessionLocal() as db:
        req = db.get(RequestItem, request_id)
        if not req:
            return
        req.match_status = "running"
        db.commit()
        db.query(MatchResult).filter(MatchResult.request_id == request_id).delete()
        db.query(ApprovedMatch).filter(ApprovedMatch.request_id == request_id, ApprovedMatch.approval_source == "auto").delete()
        items = db.scalars(select(ParsedItem).where(ParsedItem.request_id == request_id)).all()
        products = db.scalars(select(ProductCatalog).where(ProductCatalog.is_active.is_(True))).all()
        hybrid_ok = False
        if settings.hybrid_enabled and products:
            try:
                ensure_catalog_index(products)
                hybrid_ok = True
            except Exception:
                hybrid_ok = False
        stop_words = set(load_stop_words())
        match_rules = list_catalog_match_rules()
        auto, review = 0, 0
        for item in items:
            blocked_word = _contains_stop_word(item.item_name, stop_words)
            if blocked_word:
                review += 1
                db.add(
                    MatchResult(
                        request_id=request_id,
                        parsed_item_id=item.id,
                        candidate_product_id=None,
                        score=0.0,
                        status="needs_review",
                        explanation={"reasons": [f"stop_word:{blocked_word}"]},
                    )
                )
                continue
            best = None
            best_score = 0.2
            reasons: list[str] = []
            scored: list[tuple[ProductCatalog, float, list[str]]] = []
            item_sizes = _size_tokens(item.item_name)
            item_numbers = _spec_number_tokens(item.item_name)
            item_du = _du_tokens(item.item_name)
            item_fr = _fraction_tokens(item.item_name)
            vector_scores: dict[str, float] = {}
            if hybrid_ok and item.item_name:
                vector_scores = _vector_recall_scores(item.item_name)
            for p in products:
                local_reasons: list[str] = []
                allowed_by_rule, rule_reasons = apply_catalog_rule_filter(item.item_name, p.name, match_rules)
                if not allowed_by_rule:
                    continue
                local_reasons.extend(rule_reasons)
                product_du = _du_tokens(p.name)
                product_fr = _fraction_tokens(p.name)
                spec_score = 0.0
                if item.brand and item.brand.lower() == p.brand.lower():
                    spec_score += 0.08
                    local_reasons.append("brand_exact")
                name_sim = _name_similarity(item.item_name, p.name)
                lexical = 0.7 * name_sim
                vec = vector_scores.get(p.id, 0.0)
                if lexical > 0:
                    local_reasons.append(f"name_similarity:{name_sim}")
                if vec > 0:
                    local_reasons.append(f"vector_similarity:{round(vec, 3)}")
                product_sizes = _size_tokens(p.name)
                if item_sizes and product_sizes:
                    overlap = item_sizes & product_sizes
                    if overlap:
                        size_bonus = min(0.22, 0.11 * len(overlap))
                        spec_score += size_bonus
                        local_reasons.append(f"size_overlap:{','.join(sorted(overlap))}")
                product_numbers = _spec_number_tokens(p.name)
                if item_numbers:
                    matched_numbers = item_numbers & product_numbers
                    missed_numbers = item_numbers - product_numbers
                    if matched_numbers:
                        num_bonus = min(0.28, 0.08 * len(matched_numbers))
                        spec_score += num_bonus
                        local_reasons.append(f"number_overlap:{','.join(sorted(matched_numbers))}")
                    if missed_numbers:
                        num_penalty = min(0.12, 0.03 * len(missed_numbers))
                        spec_score -= num_penalty
                        local_reasons.append(f"number_miss:{','.join(sorted(missed_numbers))}")
                    # If source row contains one primary spec size (e.g., "гильза 16"),
                    # candidates without this size should be strongly demoted.
                    if len(item_numbers) == 1:
                        only_num = next(iter(item_numbers))
                        if only_num in matched_numbers:
                            spec_score += 0.2
                            local_reasons.append("single_size_exact_match")
                        else:
                            spec_score -= 0.24
                            local_reasons.append("single_size_mismatch")
                if item_du:
                    if item_du & product_du:
                        spec_score += 0.12
                        local_reasons.append("du_overlap")
                    else:
                        spec_score -= 0.06
                        local_reasons.append("du_miss")
                if item_fr:
                    if item_fr & product_fr:
                        spec_score += 0.1
                        local_reasons.append("fraction_overlap")
                    else:
                        spec_score -= 0.05
                        local_reasons.append("fraction_miss")
                penalty, penalty_reason = _semantic_penalty(item.item_name, p.name)
                if penalty > 0:
                    spec_score -= penalty
                    if penalty_reason:
                        local_reasons.append(penalty_reason)
                bonus, bonus_reason = _semantic_bonus(item.item_name, p.name)
                if bonus > 0:
                    spec_score += bonus
                    if bonus_reason:
                        local_reasons.append(bonus_reason)
                ov_bonus, ov_reason = _token_overlap_bonus(item.item_name, p.name)
                if ov_bonus > 0:
                    spec_score += ov_bonus
                    if ov_reason:
                        local_reasons.append(ov_reason)
                lexical = max(0.0, min(1.0, lexical))
                spec_score = max(-0.3, min(0.5, spec_score))
                rerank = max(0.0, min(1.0, lexical + spec_score))
                vec = max(0.0, min(1.0, vec))
                vec_weight = 0.75 if vec >= 0.35 else 0.45
                if penalty > 0:
                    vec_weight = min(vec_weight, 0.35)
                if bonus > 0 or ov_bonus > 0:
                    vec_weight = max(vec_weight, 0.55)
                score = (vec_weight * vec) + ((1 - vec_weight) * rerank)
                if score > 0:
                    scored.append((p, score, local_reasons))
                if score > best_score:
                    best, best_score, reasons = p, score, local_reasons
            top_candidates = sorted(scored, key=lambda x: x[1], reverse=True)[:5]
            second_score = top_candidates[1][1] if len(top_candidates) > 1 else 0.0
            margin = best_score - second_score
            dynamic_threshold = min(max(threshold, 0.0), 0.78)
            status = "auto_matched" if (best_score >= dynamic_threshold and margin >= 0.08) else "needs_review"
            low_confidence_cutoff = max(0.56, dynamic_threshold - 0.1)
            strong_signal = any(
                str(x).startswith("number_overlap") or str(x).startswith("semantic_bonus")
                for x in (reasons or [])
            )
            disqualifying_penalty = any(
                str(x).startswith("semantic_penalty:not_pipe_product")
                or str(x).startswith("semantic_penalty:insulation_missing")
                or str(x).startswith("semantic_penalty:tube_accessory")
                for x in (reasons or [])
            )
            unresolved = disqualifying_penalty or (best_score < low_confidence_cutoff and not strong_signal)
            if unresolved:
                best = None
                reasons = (reasons or []) + [f"low_confidence_no_replacement:{round(best_score, 3)}"]
                top_candidates = []
            if status == "auto_matched" and best:
                auto += 1
                db.add(
                    ApprovedMatch(
                        request_id=request_id,
                        parsed_item_id=item.id,
                        product_id=best.id,
                        approved_by=req.created_by,
                        approval_source="auto",
                    )
                )
            else:
                review += 1
            if not top_candidates:
                db.add(
                    MatchResult(
                        request_id=request_id,
                        parsed_item_id=item.id,
                        candidate_product_id=best.id if (best and not unresolved) else None,
                        score=min(best_score, 0.99),
                        status=status,
                        explanation={"reasons": reasons, "margin": round(margin, 4), "threshold": dynamic_threshold},
                    )
                )
            else:
                for idx, (cand, cand_score, cand_reasons) in enumerate(top_candidates):
                    db.add(
                        MatchResult(
                            request_id=request_id,
                            parsed_item_id=item.id,
                            candidate_product_id=cand.id,
                            score=min(cand_score, 0.99),
                            status=status if idx == 0 else "needs_review",
                            explanation={"reasons": cand_reasons, "margin": round(margin, 4), "threshold": dynamic_threshold},
                        )
                    )
        req.matched_items = auto
        req.needs_review_items = review
        req.match_status = "completed"
        req.status = "completed" if review == 0 else "needs_review"
        req.updated_at = datetime.now(timezone.utc)
        db.commit()


def export_task(request_id: str, export_id: str, include_unmatched: bool = True) -> None:
    with SessionLocal() as db:
        ex = db.get(ExportFile, export_id)
        if not ex:
            return
        items = db.scalars(select(ParsedItem).where(ParsedItem.request_id == request_id)).all()
        approved = db.scalars(select(ApprovedMatch).where(ApprovedMatch.request_id == request_id)).all()
        approved_by_item = {a.parsed_item_id: a for a in approved}
        product_ids = [a.product_id for a in approved if a.product_id]
        products = (
            {p.id: p for p in db.scalars(select(ProductCatalog).where(ProductCatalog.id.in_(product_ids))).all()}
            if product_ids
            else {}
        )
        filepath = export_path(export_id, ex.format)
        rows: list[dict] = []
        for item in items:
            a = approved_by_item.get(item.id)
            product = products.get(a.product_id) if a and a.product_id else None
            if not include_unmatched and not product:
                continue
            rows.append(
                {
                    "parsed_item_id": item.id,
                    "kp_item_name": item.item_name or "",
                    "kp_article": item.article or "",
                    "kp_brand": item.brand or "",
                    "kp_quantity": item.quantity,
                    "matched_product_id": product.id if product else "",
                    "matched_sku": product.sku if product else "",
                    "matched_name": product.name if product else "",
                    "matched_brand": product.brand if product else "",
                    "approval_source": a.approval_source if a else "",
                }
            )
        if ex.format == "xlsx":
            wb = Workbook()
            ws = wb.active
            ws.title = "export"
            headers = [
                "parsed_item_id",
                "kp_item_name",
                "kp_article",
                "kp_brand",
                "kp_quantity",
                "matched_product_id",
                "matched_sku",
                "matched_name",
                "matched_brand",
                "approval_source",
            ]
            ws.append(headers)
            for row in rows:
                ws.append([row[h] for h in headers])
            wb.save(filepath)
        else:
            lines = [
                "parsed_item_id,kp_item_name,kp_article,kp_brand,kp_quantity,matched_product_id,matched_sku,matched_name,matched_brand,approval_source"
            ]
            for row in rows:
                values = [
                    row["parsed_item_id"],
                    str(row["kp_item_name"]).replace(",", " "),
                    str(row["kp_article"]).replace(",", " "),
                    str(row["kp_brand"]).replace(",", " "),
                    str(row["kp_quantity"]),
                    row["matched_product_id"],
                    str(row["matched_sku"]).replace(",", " "),
                    str(row["matched_name"]).replace(",", " "),
                    str(row["matched_brand"]).replace(",", " "),
                    row["approval_source"],
                ]
                lines.append(",".join(values))
            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
        ex.storage_key = filepath
        ex.status = "completed"
        db.commit()
